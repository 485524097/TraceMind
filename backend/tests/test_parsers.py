from pathlib import Path
from typing import cast

import pytest
from docx import Document
from pypdf import PdfReader as RealPdfReader

from app.parsing.base import ParseContext
from app.parsing.code import LANGUAGES, CodeParser
from app.parsing.docx import DocxParser
from app.parsing.exceptions import (
    DocumentEncodingError,
    DocumentParseError,
    NoExtractableTextError,
    ParseLimitExceededError,
    PdfEncryptedError,
)
from app.parsing.java import JavaTreeSitterParser
from app.parsing.markdown import MarkdownParser
from app.parsing.pdf import PdfParser
from app.parsing.registry import ParserRegistry
from app.parsing.text import PlainTextParser

CONTEXT = ParseContext(max_extracted_chars=10_000, max_pdf_pages=10)


def write(path: Path, content: bytes) -> Path:
    path.write_bytes(content)
    return path


def test_plain_text_utf8_sig_crlf_chinese_and_lines(tmp_path: Path) -> None:
    path = write(tmp_path / "sample.txt", "first\r\n\r\n中文".encode("utf-8-sig"))
    parsed = PlainTextParser().parse(path, CONTEXT)

    assert [block.text for block in parsed.blocks] == ["first", "中文"]
    assert [(block.start_line, block.end_line) for block in parsed.blocks] == [(1, 1), (3, 3)]


@pytest.mark.parametrize("content", [b"\xff", b"text\x00binary"])
def test_plain_text_rejects_invalid_utf8_and_nul(tmp_path: Path, content: bytes) -> None:
    with pytest.raises(DocumentEncodingError):
        PlainTextParser().parse(write(tmp_path / "sample.txt", content), CONTEXT)


def test_plain_text_rejects_empty_and_character_limit(tmp_path: Path) -> None:
    with pytest.raises(NoExtractableTextError):
        PlainTextParser().parse(write(tmp_path / "empty.txt", b" \n"), CONTEXT)
    with pytest.raises(ParseLimitExceededError):
        PlainTextParser().parse(
            write(tmp_path / "large.txt", b"12345"),
            ParseContext(max_extracted_chars=4, max_pdf_pages=10),
        )


def test_markdown_headings_sections_fences_and_lines(tmp_path: Path) -> None:
    content = "# 安装\n说明\n\n```python\nprint('ok')\n```\n## 配置\n值"
    parsed = MarkdownParser().parse(write(tmp_path / "sample.md", content.encode()), CONTEXT)

    assert [block.block_type for block in parsed.blocks] == [
        "heading",
        "paragraph",
        "code",
        "heading",
        "paragraph",
    ]
    assert parsed.blocks[1].section_title == "安装"
    assert parsed.blocks[2].language == "python"
    assert (parsed.blocks[2].start_line, parsed.blocks[2].end_line) == (4, 6)
    assert parsed.blocks[-1].section_title == "配置"


def test_markdown_bare_backtick_fence_preserves_content_and_metadata(tmp_path: Path) -> None:
    content = "# 简历\n\n项目经历\n\n```\nTraceMind\nFastAPI\nCelery\n```"

    parsed = MarkdownParser().parse(write(tmp_path / "resume.md", content.encode()), CONTEXT)
    code = next(block for block in parsed.blocks if block.block_type == "code")

    assert code.language is None
    assert code.block_type == "code"
    assert code.section_title == "简历"
    assert (code.start_line, code.end_line) == (5, 9)
    assert code.text == "```\nTraceMind\nFastAPI\nCelery\n```"


def test_markdown_bare_tilde_fence_preserves_content_and_lines(tmp_path: Path) -> None:
    content = "~~~\nplain content\n~~~"

    parsed = MarkdownParser().parse(write(tmp_path / "plain.md", content.encode()), CONTEXT)
    code = parsed.blocks[0]

    assert code.block_type == "code"
    assert code.language is None
    assert code.text == content
    assert (code.start_line, code.end_line) == (1, 3)


@pytest.mark.parametrize("language", ["python", "java", "javascript"])
def test_markdown_language_fence_preserves_language(tmp_path: Path, language: str) -> None:
    content = f'```{language}\nprint("ok")\n```'

    parsed = MarkdownParser().parse(write(tmp_path / "language.md", content.encode()), CONTEXT)

    assert parsed.blocks[0].language == language


def test_markdown_fence_uses_first_info_word_as_language(tmp_path: Path) -> None:
    content = '```python linenums\nprint("ok")\n```'

    parsed = MarkdownParser().parse(write(tmp_path / "info.md", content.encode()), CONTEXT)

    assert parsed.blocks[0].language == "python"


def test_markdown_resume_style_document_with_bare_fence(tmp_path: Path) -> None:
    content = "\n".join(
        [
            "# 候选人概览",
            "",
            "## 项目经历",
            "",
            "- **项目名称**：[TraceMind](https://example.com)",
            "- 技术栈：Python、FastAPI、Celery",
            "",
            "---",
            "",
            "```",
            "uv run pytest",
            "```",
            "",
            "具备中文技术文档整理与问题分析经验。",
        ]
    )

    parsed = MarkdownParser().parse(
        write(tmp_path / "synthetic-resume.md", content.encode()), CONTEXT
    )
    code = next(block for block in parsed.blocks if block.block_type == "code")

    assert {block.block_type for block in parsed.blocks} >= {"heading", "paragraph", "code"}
    assert code.language is None
    assert code.section_title == "项目经历"
    assert (code.start_line, code.end_line) == (10, 12)
    assert code.text == "```\nuv run pytest\n```"


@pytest.mark.parametrize(("extension", "language"), sorted(LANGUAGES.items()))
def test_code_parser_language_indentation_and_lines(
    tmp_path: Path, extension: str, language: str
) -> None:
    path = write(tmp_path / f"sample{extension}", b"class A:\n    value = 1\n\nreturn value")
    parsed = CodeParser().parse(path, CONTEXT)

    assert parsed.blocks[0].text == "class A:\n    value = 1"
    assert parsed.blocks[0].language == language
    assert (parsed.blocks[0].start_line, parsed.blocks[0].end_line) == (1, 2)
    assert (parsed.blocks[1].start_line, parsed.blocks[1].end_line) == (4, 4)


class FakePage:
    def __init__(self, text: str | None) -> None:
        self.text = text

    def extract_text(self) -> str | None:
        return self.text


class FakeReader:
    def __init__(self, pages: list[FakePage], *, encrypted: bool = False) -> None:
        self.pages = pages
        self.is_encrypted = encrypted

    def decrypt(self, _password: str) -> int:
        return 0


def test_pdf_pages_are_one_based_and_blank_pages_are_skipped(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(
        "app.parsing.pdf.PdfReader",
        lambda *_args, **_kwargs: FakeReader([FakePage("one"), FakePage(None), FakePage("三")]),
    )
    parsed = PdfParser().parse(tmp_path / "sample.pdf", CONTEXT)
    assert [(block.page_number, block.text) for block in parsed.blocks] == [(1, "one"), (3, "三")]


def test_pdf_encryption_no_text_page_and_character_limits(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(
        "app.parsing.pdf.PdfReader", lambda *_args, **_kwargs: FakeReader([], encrypted=True)
    )
    with pytest.raises(PdfEncryptedError):
        PdfParser().parse(tmp_path / "encrypted.pdf", CONTEXT)

    monkeypatch.setattr(
        "app.parsing.pdf.PdfReader", lambda *_args, **_kwargs: FakeReader([FakePage(None)])
    )
    with pytest.raises(NoExtractableTextError):
        PdfParser().parse(tmp_path / "blank.pdf", CONTEXT)

    monkeypatch.setattr(
        "app.parsing.pdf.PdfReader", lambda *_args, **_kwargs: FakeReader([FakePage("12345")])
    )
    with pytest.raises(ParseLimitExceededError):
        PdfParser().parse(
            tmp_path / "large.pdf", ParseContext(max_extracted_chars=4, max_pdf_pages=10)
        )


def test_pdf_page_limit_and_invalid_document(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(
        "app.parsing.pdf.PdfReader",
        lambda *_args, **_kwargs: FakeReader([FakePage("one"), FakePage("two")]),
    )
    with pytest.raises(ParseLimitExceededError):
        PdfParser().parse(
            tmp_path / "pages.pdf", ParseContext(max_extracted_chars=100, max_pdf_pages=1)
        )
    monkeypatch.setattr("app.parsing.pdf.PdfReader", RealPdfReader)
    with pytest.raises(DocumentParseError):
        PdfParser().parse(write(tmp_path / "broken.pdf", b"not a pdf"), CONTEXT)


def test_docx_preserves_heading_paragraph_table_order_and_chinese(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("架构", level=1)
    document.add_paragraph("")
    document.add_paragraph("中文说明")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "键"
    table.cell(0, 1).text = "值"
    document.save(path)

    parsed = DocxParser().parse(path, CONTEXT)

    assert [block.block_type for block in parsed.blocks] == ["heading", "paragraph", "table"]
    assert [block.section_title for block in parsed.blocks] == ["架构", "架构", "架构"]
    assert parsed.blocks[-1].text == "键\t值"
    assert parsed.blocks[-1].page_number is None


def test_docx_rejects_broken_document(tmp_path: Path) -> None:
    with pytest.raises(DocumentParseError):
        DocxParser().parse(write(tmp_path / "broken.docx", b"not a zip"), CONTEXT)


def test_registry_maps_every_supported_extension() -> None:
    registry = ParserRegistry()
    expected = {
        ".md",
        ".txt",
        ".json",
        ".yaml",
        ".yml",
        ".xml",
        ".properties",
        ".java",
        ".jsp",
        ".js",
        ".ts",
        ".vue",
        ".sql",
        ".py",
        ".pdf",
        ".docx",
    }
    assert registry.supported_extensions == expected
    assert cast(object, registry.get(".MD")).parser_name == "markdown"
    assert isinstance(registry.get(".java"), JavaTreeSitterParser)
    assert all(isinstance(registry.get(extension), CodeParser) for extension in LANGUAGES)


def parse_java(tmp_path: Path, source: str):
    return JavaTreeSitterParser().parse(
        write(tmp_path / "Sample.java", source.encode("utf-8")),
        CONTEXT,
    )


def symbol_blocks(parsed):
    return [block for block in parsed.blocks if block.symbol_kind is not None]


def test_java_types_members_nested_and_initializers(tmp_path: Path) -> None:
    source = """package demo;
public class Outer {
  private String first, second;
  static { boot(); }
  { init(); }
  Outer() {}
  void work() { if (true) { boot(); } }
  interface Nested { void call(); }
}
interface Contract {}
enum State { READY, DONE }
record Item(String name) { Item {} }
@interface Marker {}
"""
    parsed = parse_java(tmp_path, source)
    symbols = symbol_blocks(parsed)

    assert {"Outer", "Nested", "Contract", "State", "Item", "Marker"} <= {
        block.symbol_name for block in symbols if block.symbol_kind == "type"
    }
    assert [(block.symbol_kind, block.symbol_name) for block in symbols].count(
        ("initializer", "<init-block>")
    ) == 1
    assert ("initializer", "<clinit>") in [
        (block.symbol_kind, block.symbol_name) for block in symbols
    ]
    assert ("constructor", "Outer") in [(block.symbol_kind, block.symbol_name) for block in symbols]
    assert ("constructor", "Item") in [(block.symbol_kind, block.symbol_name) for block in symbols]
    assert {"READY", "DONE"} <= {
        block.symbol_name for block in symbols if block.symbol_kind == "enum_constant"
    }
    field = next(block for block in symbols if block.symbol_kind == "field")
    assert field.symbol_name == "first"
    assert field.symbol_signature == "private String first, second;"
    outer = next(block for block in symbols if block.symbol_name == "Outer")
    assert outer.text.endswith("{")
    assert "private String" not in outer.text
    assert next(
        block for block in symbols if block.symbol_name == "call"
    ).symbol_qualified_name == ("demo.Outer.Nested.call")


def test_java_overloads_annotations_generics_throws_and_qualified_names(
    tmp_path: Path,
) -> None:
    parsed = parse_java(
        tmp_path,
        """package demo;
class UserService {
  @Deprecated public <T> T source(String username) throws Exception { return null; }
  long source(long id) { return id; }
}
""",
    )
    methods = [block for block in symbol_blocks(parsed) if block.symbol_kind == "method"]

    assert [block.symbol_name for block in methods] == ["source", "source"]
    assert {block.symbol_qualified_name for block in methods} == {"demo.UserService.source"}
    assert len({block.symbol_signature for block in methods}) == 2
    assert "@Deprecated public <T> T source(String username) throws Exception" in (
        methods[0].symbol_signature or ""
    )


def test_java_javadoc_attachment_and_uncovered_source(tmp_path: Path) -> None:
    source = """// header
package demo;
import java.util.List;
/* ordinary */
/** type docs */
class Sample {
  /** field docs */
  int value;
  /** detached */
  // blocker
  void run() {}
}
"""
    parsed = parse_java(tmp_path, source)
    sample = next(block for block in symbol_blocks(parsed) if block.symbol_name == "Sample")
    field = next(block for block in symbol_blocks(parsed) if block.symbol_kind == "field")
    method = next(block for block in symbol_blocks(parsed) if block.symbol_kind == "method")
    ordinary = "\n".join(block.text for block in parsed.blocks if block.symbol_kind is None)

    assert sample.text.startswith("/** type docs */")
    assert field.text.startswith("/** field docs */")
    assert not method.text.startswith("/**")
    assert "/** type docs */" not in ordinary
    assert "/** field docs */" not in ordinary
    assert "/** detached */" in ordinary
    assert "// header" in ordinary and "package demo;" in ordinary
    assert "import java.util.List;" in ordinary and "/* ordinary */" in ordinary


def test_java_unicode_crlf_partial_error_and_stable_output(tmp_path: Path) -> None:
    source = (
        "package 示例;\r\n"
        "class 用户 {\r\n"
        '  String 名称() { return "中"; }\r\n'
        "  void broken( {\r\n"
        "}\r\n"
    )
    first = parse_java(tmp_path, source)
    second = parse_java(tmp_path, source)
    method = next(block for block in symbol_blocks(first) if block.symbol_name == "名称")

    assert "java_partial_syntax_tree" in first.warnings
    assert method.symbol_qualified_name == "示例.用户.名称"
    assert (method.start_line, method.end_line) == (3, 3)
    assert first.blocks == second.blocks


def test_java_parser_falls_back_when_language_is_unavailable(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr("app.parsing.java.JAVA_LANGUAGE", None)
    parsed = parse_java(tmp_path, "class Sample { void run() {} }")

    assert parsed.parser_name == "code"
    assert parsed.warnings == ["java_parser_fallback"]
    assert all(block.symbol_kind is None for block in parsed.blocks)
    assert all(block.symbol_lookup_keys is None for block in parsed.blocks)


def test_java_lookup_keys_cover_types_members_fields_initializers_and_record(
    tmp_path: Path,
) -> None:
    parsed = parse_java(
        tmp_path,
        """package demo;
record Item(String name, int count) { Item {} }
class Outer {
  private int first, second;
  static { boot(); }
  { init(); }
  Outer(String value) {}
  void run() {}
  class Nested { void call(int[] ids) {} }
}
enum State { READY, DONE }
""",
    )
    symbols = symbol_blocks(parsed)

    outer = next(
        block for block in symbols if block.symbol_kind == "type" and block.symbol_name == "Outer"
    )
    assert outer.symbol_lookup_keys == ["v1:type:demo.Outer", "v1:type:Outer"]
    nested = next(
        block for block in symbols if block.symbol_kind == "type" and block.symbol_name == "Nested"
    )
    assert nested.symbol_lookup_keys == [
        "v1:type:demo.Outer.Nested",
        "v1:type:Outer.Nested",
        "v1:type:Nested",
    ]
    nested_method = next(
        block for block in symbols if block.symbol_kind == "method" and block.symbol_name == "call"
    )
    assert "v1:method:Outer.Nested#call(int[])" in (nested_method.symbol_lookup_keys or [])
    field = next(block for block in symbols if block.symbol_kind == "field")
    assert field.symbol_name == "first"
    assert field.symbol_lookup_keys == [
        "v1:field:demo.Outer#first",
        "v1:field:demo.Outer#second",
        "v1:field:Outer#first",
        "v1:field:Outer#second",
    ]
    constructor = next(
        block
        for block in symbols
        if block.symbol_kind == "constructor" and block.symbol_name == "Outer"
    )
    assert "v1:constructor:demo.Outer#<init>(String)" in (constructor.symbol_lookup_keys or [])
    compact = next(
        block
        for block in symbols
        if block.symbol_kind == "constructor" and block.symbol_name == "Item"
    )
    assert "v1:constructor:demo.Item#<init>(String,int)" in (compact.symbol_lookup_keys or [])
    assert next(
        block for block in symbols if block.symbol_name == "<clinit>"
    ).symbol_lookup_keys == [
        "v1:initializer:demo.Outer#<clinit>",
        "v1:initializer:Outer#<clinit>",
    ]
    assert next(
        block for block in symbols if block.symbol_name == "<init-block>"
    ).symbol_lookup_keys == [
        "v1:initializer:demo.Outer#<init-block>",
        "v1:initializer:Outer#<init-block>",
    ]
    assert "v1:enum_constant:demo.State#READY" in (
        next(block for block in symbols if block.symbol_name == "READY").symbol_lookup_keys or []
    )
    assert all(
        block.symbol_lookup_keys
        and len(block.symbol_lookup_keys) == len(set(block.symbol_lookup_keys))
        for block in symbols
    )


def test_java_lookup_keys_distinguish_overloads_and_normalize_java_types(tmp_path: Path) -> None:
    source = """package de\u0301mo;
class UserService {
  void source(Map < String, ? extends User > users, String... names) {}
  void source(String values[]) {}
}
"""
    parsed = parse_java(tmp_path, source)
    methods = [block for block in symbol_blocks(parsed) if block.symbol_kind == "method"]

    assert methods[0].symbol_lookup_keys == [
        "v1:method:démo.UserService#source",
        "v1:method:démo.UserService#source(Map<String,? extends User>,String[])",
        "v1:method:UserService#source",
        "v1:method:UserService#source(Map<String,? extends User>,String[])",
    ]
    assert "v1:method:démo.UserService#source(String[])" in (methods[1].symbol_lookup_keys or [])
    assert methods[0].symbol_lookup_keys != methods[1].symbol_lookup_keys
    second = parse_java(tmp_path, source)
    assert parsed.blocks == second.blocks
